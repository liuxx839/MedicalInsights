#main.py
import streamlit as st
from layout import setup_layout
from functions import (
    setup_client, generate_tag, generate_diseases_tag, rewrite,
    prob_identy, generate_structure_data
)
from config import (
    topics, diseases, institutions, departments, persons,
    primary_topics_list, primary_diseases_list, colors
)
from streamlit_extras.stylable_container import stylable_container
import pandas as pd
import json
import time
import ast
import io
from openai import OpenAI
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from dagrelation import DAGRelations
from datadescription import DataDescription
import numpy as np
from datetime import datetime
from prophet import Prophet
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import base64
import re
import os

# 新增的导入（原来没有）
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.cluster import KMeans, DBSCAN
from sklearn.preprocessing import StandardScaler
from sklearn.metrics import silhouette_score
from sklearn.manifold import TSNE
from sklearn.neighbors import NearestNeighbors
from scipy.stats import chi2_contingency
import matplotlib as mpl

# Set Matplotlib font configuration
plt.rcParams['font.family'] = 'sans-serif'
plt.rcParams['font.sans-serif'] = ['Noto Sans CJK JP', 'WenQuanYi Zen Hei', 'sans-serif']
plt.rcParams['axes.unicode_minus'] = False


model_choice_research, client_research = setup_client(model_choice = 'gemini-2.0-flash')
 
def create_mermaid_html_from_edges(dag_edges):
    """
    Create a Mermaid flowchart HTML from a list of edges.
    
    Parameters:
    dag_edges (list): List of tuples representing edges. Each tuple can be:
                     - ('var1', 'var2') for a single edge
                     - (['var1', 'var2'], 'var3') for multiple sources to one target
    
    Returns:
    str: HTML code with embedded Mermaid flowchart
    """
    mermaid_code = ['flowchart TD']
    
    for edge in dag_edges:
        if isinstance(edge[0], list):
            # Handle multiple sources to one target
            source_vars = edge[0]
            target_var = edge[1]
            
            # Create individual connections for each source to target
            for source_var in source_vars:
                mermaid_code.append(f'    {source_var} --> {target_var}')
        else:
            # Simple one-to-one edge
            source_var = edge[0]
            target_var = edge[1]
            mermaid_code.append(f'    {source_var} --> {target_var}')
    
    # Join all lines with newlines
    mermaid_code_str = '\n'.join(mermaid_code)
    
    # Create the HTML with the embedded Mermaid code and zoom controls
    mermaid_html = f"""
<div id="mermaid-container" style="width:100%; height:100%; position:relative;">
    <div class="mermaid" id="mermaid-diagram">
    {mermaid_code_str}
    </div>
    
    <div style="position:absolute; top:10px; right:10px; background:white; padding:5px; border:1px solid #ccc; border-radius:5px;">
        <button onclick="zoomIn()" style="margin-right:5px;">➕</button>
        <button onclick="zoomOut()">➖</button>
        <button onclick="resetZoom()" style="margin-left:5px;">🔄</button>
    </div>
</div>

<script type="module">
    import mermaid from 'https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.esm.min.mjs';
    mermaid.initialize({{ startOnLoad: true }});
    
    window.currentZoom = 1;
    
    window.zoomIn = function() {{
        window.currentZoom += 0.1;
        document.getElementById('mermaid-diagram').style.transform = `scale(${{window.currentZoom}})`;
        document.getElementById('mermaid-diagram').style.transformOrigin = 'top left';
    }};
    
    window.zoomOut = function() {{
        if (window.currentZoom > 0.5) {{
            window.currentZoom -= 0.1;
            document.getElementById('mermaid-diagram').style.transform = `scale(${{window.currentZoom}})`;
            document.getElementById('mermaid-diagram').style.transformOrigin = 'top left';
        }}
    }};
    
    window.resetZoom = function() {{
        window.currentZoom = 1;
        document.getElementById('mermaid-diagram').style.transform = 'scale(1)';
    }};
</script>
"""
    
    return mermaid_html
    
def create_word_document(qa_response, fact_check_result=None):
    """
    创建包含QA响应和事实核查结果的Word文档，支持Markdown标题转换为Word样式
    """
    doc = Document()
    
    # 设置标题
    title = doc.add_heading('Medical Knowledge Base Q&A Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加时间戳
    timestamp = doc.add_paragraph()
    timestamp.add_run(f"生成时间: {time.strftime('%Y-%m-%d %H:%M:%S')}").italic = True
    
    # 添加QA响应
    doc.add_heading('回应内容', level=1)
    
    # 处理QA响应中的Markdown格式
    lines = qa_response.split('\n')
    current_text = []
    
    for line in lines:
        # 检查是否是标题行
        if line.strip().startswith('#'):
            # 如果之前有累积的文本，先添加为段落
            if current_text:
                doc.add_paragraph(''.join(current_text))
                current_text = []
            
            # 计算标题级别
            level = 1
            line = line.strip()
            while line.startswith('#'):
                level += 1
                line = line[1:]
            level = min(level, 9)  # Word支持最多9级标题
            
            # 移除可能的引用标记 [1,2,3] 并保存
            citation = ''
            if '[' in line and ']' in line:
                main_text = line[:line.find('[')].strip()
                citation = line[line.find('['):].strip()
                heading = doc.add_heading(main_text, level=level-1)
                if citation:
                    heading.add_run(f" {citation}").italic = True
            else:
                doc.add_heading(line.strip(), level=level-1)
        else:
            current_text.append(line + '\n')
    
    # 添加最后剩余的文本
    if current_text:
        doc.add_paragraph(''.join(current_text))
    
    # 如果有事实核查结果，添加到文档
    if fact_check_result:
        doc.add_heading('事实核查结果', level=1)
        # 对事实核查结果也进行相同的处理
        lines = fact_check_result.split('\n')
        current_text = []
        
        for line in lines:
            if line.strip().startswith('#'):
                if current_text:
                    doc.add_paragraph(''.join(current_text))
                    current_text = []
                
                level = 1
                line = line.strip()
                while line.startswith('#'):
                    level += 1
                    line = line[1:]
                level = min(level, 9)
                
                if '[' in line and ']' in line:
                    main_text = line[:line.find('[')].strip()
                    citation = line[line.find('['):].strip()
                    heading = doc.add_heading(main_text, level=level-1)
                    if citation:
                        heading.add_run(f" {citation}").italic = True
                else:
                    doc.add_heading(line.strip(), level=level-1)
            else:
                current_text.append(line + '\n')
        
        if current_text:
            doc.add_paragraph(''.join(current_text))
    
    # 保存到内存中
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

def extract_dag_edges(text_content):
    # 从文本中提取 dag_edges 部分
    if "dag_edges = [" in text_content:
        start_idx = text_content.find("dag_edges = [")
        start_idx = text_content.find("[", start_idx)
        
        # 找到匹配的结束括号
        open_brackets = 1
        end_idx = start_idx + 1
        
        while open_brackets > 0 and end_idx < len(text_content):
            if text_content[end_idx] == '[':
                open_brackets += 1
            elif text_content[end_idx] == ']':
                open_brackets -= 1
            end_idx += 1
        
        # 提取 dag_edges 列表内容
        dag_edges_str = text_content[start_idx:end_idx]
        
        # 使用 ast.literal_eval 安全地将字符串转换为 Python 对象
        try:
            dag_edges = ast.literal_eval(dag_edges_str)
            return dag_edges
        except (SyntaxError, ValueError) as e:
            print(f"解析错误: {e}")
            return None
    
    return None

def setup_spreadsheet_analysis():
    # st.markdown(
    #     """
    # <h1 style='text-align: center; font-size: 18px; font-weight: bold;'>Spreadsheet Analysis</h1>
    # <h6 style='text-align: center; font-size: 12px;'>上传Excel/CSV文件或粘贴JSON数据进行分析</h6>
    # <br><br><br>
    # """,
    #     unsafe_allow_html=True,
    # )
    
    # Initialize all session state variables
    if "analysis_response" not in st.session_state:
        st.session_state.analysis_response = ""
    if "analysis_reasoning" not in st.session_state:
        st.session_state.analysis_reasoning = ""
    if "dag_edges" not in st.session_state:
        st.session_state.dag_edges = ""
    if "sample_data" not in st.session_state:
        st.session_state.sample_data = {}
    if "business_report" not in st.session_state:
        st.session_state.business_report = ""
    if "dag_report" not in st.session_state:
        st.session_state.dag_report = ""
    if "dag_reasoning" not in st.session_state:
        st.session_state.dag_reasoning = ""
    if "df" not in st.session_state:
        st.session_state.df = None
    
    # Create tabs for different input methods
    tab1, tab2 = st.tabs(["文件上传", "JSON粘贴"])
    
    with tab1:
        uploaded_file = st.file_uploader("上传文件(xlsx/csv)", type=['xlsx', 'csv'])
        
        if uploaded_file is not None:
            try:
                # 根据文件类型读取数据
                if uploaded_file.name.endswith('.csv'):
                    df = pd.read_csv(uploaded_file)
                else:
                    df = pd.read_excel(uploaded_file)
                    
                # 处理列名，将空格和特殊符号替换为下划线
                df.columns = [re.sub(r'[^\w]', '_', col) for col in df.columns]
                
                # 保存DataFrame到session state
                st.session_state.df = df
                
                # 显示前10行数据
                st.write("数据预览:")
                st.dataframe(df)
                
                # 保存sample数据到session state
                st.session_state.sample_data = df.sample(10).to_dict()
                
            except Exception as e:
                st.error(f"处理文件时出错: {str(e)}")
    
    with tab2:
        json_input = st.text_area("粘贴JSON数据:", height=200)
        process_json_button = st.button("处理JSON数据")
        
        if process_json_button and json_input:
            try:
                # 清理JSON字符串 - 如果前后有额外的引号，去掉它们
                cleaned_json = json_input.strip()
                if cleaned_json.startswith('"') and cleaned_json.endswith('"'):
                    # 如果JSON被额外的引号包围，去掉这些引号并处理转义字符
                    cleaned_json = cleaned_json[1:-1].replace('\\"', '"')
                
                # 解析JSON数据
                try:
                    json_data = json.loads(cleaned_json)
                    df = pd.DataFrame(json_data)
                except Exception as e:
                    # 如果第一次尝试失败，尝试以Records格式解析
                    try:
                        json_data = json.loads(cleaned_json)
                        if isinstance(json_data, list):
                            df = pd.DataFrame(json_data)
                        else:
                            df = pd.DataFrame([json_data])
                    except:
                        st.error(f"无法解析JSON数据: {str(e)}")
                        return
                
                # 保存DataFrame到session state
                st.session_state.df = df
                
                # 显示前10行数据
                st.write("数据预览:")
                st.dataframe(df)
                
                # 保存sample数据到session state
                st.session_state.sample_data = df.head(10).to_dict()
                
            except Exception as e:
                st.error(f"处理JSON数据时出错: {str(e)}")
    
    # 只有当DataFrame可用时才显示下面的控件
    if st.session_state.df is not None:
        # 用户输入需求
        user_question = st.text_area("请输入您的需求:", height=100)
        
        # 创建两列布局，一列放生成按钮，一列放DAG分析按钮
        col1, col2 = st.columns(2)
        
        with col1:
            # 根据是否已经进行过DAG分析来显示不同的按钮文本
            button_text = "再次生成" if "business_report" in st.session_state and st.session_state.business_report else "生成"
            generate_button = st.button(button_text)
        
        with col2:
            with stylable_container(
                "dag_analysis_button",
                css_styles="""
                    button {
                        background-color: #FFA500;
                        color: white;
                    }""",
            ):
                dag_analysis_button = st.button("📊 DAG分析")
        
        # 确保有DataFrame可用于分析
        df = st.session_state.df
        
        # 处理生成按钮点击
        if generate_button and user_question:
            with st.spinner("正在分析..."):
                # 根据是否已经有business_report来决定使用哪个提示
                if "business_report" in st.session_state and st.session_state.business_report:
                    # 使用已有的business_report作为输入，重新思考DAG结构
                    response = client_research.chat.completions.create(
                        model=model_choice_research,
                        messages=[
                            {
                                "role": "system",
                                "content": """你是一个数据分析专家。请基于提供的sample数据和已有的分析报告，重新思考并优化DAG关系。
                                关注以下几点：
                                1. 仔细分析已有报告中的发现和关系
                                2. 重新评估可能的因果关系
                                3. 构建更优化的DAG边
                                4. 支持多对一的关系
                                5. 确保使用的是原始的列名，不要做任何修改
                                6. 考虑已有分析中可能被忽略的关系
                                7. 不要增加任何不存在的列名

                                请先提供详细的推理过程，然后再给出优化后的DAG定义。

                                最终输出格式必须如下：
                                ##推理过程
                                [详细的分析推理过程，包括对已有分析的评估]

                                ##定义DAG边（支持多对一关系）
                                请严格遵循下面的格式，仔细核对，务必保证使用原始列名，以及正确的括号
                                dag_edges = [
                                    ('var1', 'var2'),
                                    ('var3', 'var4'),
                                    # 多对一关系示例
                                    (['var1', 'var2'], 'var3'),
                                    (['var1', 'var2', 'var3'], 'var4')
                                ]
                                
                                ##分析说明：
                                [这里是对优化后DAG结构的解释说明，以及与之前分析的比较]
                                """
                            },
                            {
                                "role": "user",
                                "content": f"Sample数据：\n{st.session_state.sample_data}\n\n用户需求：{user_question}\n\n已有分析报告：\n{st.session_state.business_report}"
                            }
                        ],
                        temperature=0.7,
                        max_tokens=2000,
                        stream=True
                    )
                else:
                    # 原有的提示，用于首次生成
                    response = client_research.chat.completions.create(
                        model=model_choice_research,
                        messages=[
                            {
                                "role": "system", 
                                "content": """
                                你是一个数据分析专家。请基于提供的sample数据分析用户需求，并构建潜在的DAG关系。
                                关注以下几点：
                                1. 仔细分析数据列之间的关系
                                2. 识别可能的因果关系
                                3. 构建合适的DAG边
                                4. 支持多对一的关系
                                5. 确保使用的是原始的列名，不要做任何修改
                                6. 不要增加任何不存在的列名

                                最终输出格式必须如下：

                                ## 推理过程
                                [详细的分析推理过程]

                                ## 定义DAG边（支持多对一关系）
                                请严格遵循下面的格式，仔细核对，务必保证使用原始列名，以及正确的括号
                                dag_edges = [
                                    ('var1', 'var2'),
                                    ('var3', 'var4'),
                                    # 多对一关系示例
                                    (['var1', 'var2'], 'var3'),
                                    (['var1', 'var2', 'var3'], 'var4')
                                ]


                                ## 分析说明：
                                [这里是对DAG结构的解释说明，包含原则遵循情况分析]
                                """
                            },
                            {
                                "role": "user", 
                                "content": f"Sample数据：\n{st.session_state.sample_data}\n\n用户需求：{user_question}"
                            }
                        ],
                        temperature=0.7,
                        max_tokens=2000,
                        stream=True
                    )
                
                full_response = ""
                reasoning_content = ""
                
                # 创建进度显示的容器
                progress_container = st.empty()
                
                # 处理流式响应
                for chunk in response:
                    if hasattr(chunk.choices[0].delta, 'reasoning_content') and chunk.choices[0].delta.reasoning_content:
                        reasoning_content += chunk.choices[0].delta.reasoning_content
                        progress_container.markdown(f"思考过程：\n{reasoning_content}\n\n回答：\n{full_response}")
                    elif hasattr(chunk.choices[0].delta, 'content') and chunk.choices[0].delta.content:
                        full_response += chunk.choices[0].delta.content
                        progress_container.markdown(f"思考过程：\n{reasoning_content}\n\n回答：\n{full_response}")
                
                # 仅提取DAG定义部分
                dag_definition = ""
                if "dag_edges = [" in full_response:
                    dag_start = full_response.find("dag_edges = [")
                    
                    # Find the matching closing bracket by counting opening and closing brackets
                    open_brackets = 1
                    dag_end = dag_start + len("dag_edges = [")
                    
                    while open_brackets > 0 and dag_end < len(full_response):
                        if full_response[dag_end] == '[':
                            open_brackets += 1
                        elif full_response[dag_end] == ']':
                            open_brackets -= 1
                        dag_end += 1
                    
                    dag_definition = full_response[dag_start:dag_end]
                
                # 保存到session state
                st.session_state.analysis_response = full_response
                st.session_state.analysis_reasoning = reasoning_content
                st.session_state.dag_edges = dag_definition
                
                # 清空进度容器
                progress_container.empty()
        
        # 处理DAG分析按钮点击
        if dag_analysis_button and st.session_state.dag_edges:
            with st.spinner("正在进行DAG分析..."):
                try:
                    # 执行DAG分析
                    dag_edges_text = st.session_state.dag_edges
                    dag_edges = extract_dag_edges(dag_edges_text)
                    
                    full_response = ""
                    reasoning_content = ""

                    # 创建进度显示的容器
                    dag_progress = st.empty()
                    
                    if dag_edges:
                        # 执行DAG分析
                        analyzer = DAGRelations(df, dag_edges)
                        dag_report = analyzer.analyze_relations().print_report()
                        st.session_state.dag_report = dag_report

                        # 添加数据描述分析
                        data_analyzer = DataDescription(df, include_histogram=False, string_threshold=10)
                        data_analyzer.analyze_data()
                        json_output = data_analyzer.to_json()
                        st.session_state.data_description = json_output

                        # 生成商业报告
                        response = client_research.chat.completions.create(
                            model=model_choice_research,
                            messages=[
                                {
                                    "role": "system",
                                    "content": """你是一位专业的数据分析师。请基于提供的初始分析结果、DAG分析报告和数据描述信息生成一份结构化的分析报告。如果DAG分析报告不存在，请在报告中说明这一情况。

                                        报告结构应包含：

                                        # 执行摘要
                                        简明扼要地总结关键发现和建议（不超过200字）
                                        
                                        # 关键洞察
                                        ## 主要发现
                                        - 发现1: [简明描述] - 支持数据: [相关统计结果]
                                        - 发现2: [简明描述] - 支持数据: [相关统计结果]
                                        - 发现n: ...
                                        
                                        # 数据概览
                                        ## 数据基本情况
                                        以表格形式呈现：
                                        | 数据维度 | 值 |
                                        | --- | --- |
                                        | 总记录数 | X |
                                        | 变量数量 | X |
                                        | 数值型变量 | X个 (列出名称) |
                                        | 分类型变量 | X个 (列出名称) |
                                        | 缺失值情况 | 总体百分比及主要缺失列 |

                                        ## 关键变量分析
                                        针对重要变量以表格形式呈现：

                                        **数值型变量统计**
                                        | 变量名 | 均值 | 中位数 | 标准差 | 最小值 | 最大值 | 异常值比例 |
                                        | --- | --- | --- | --- | --- | --- | --- |

                                        **分类型变量统计**
                                        | 变量名 | 唯一值数量 | 最常见类别(占比) | 熵值(归一化) | 分布均衡度 |
                                        | --- | --- | --- | --- | --- | --- |

                                        # 变量关系分析
                                        ## DAG关系概述
                                        以表格形式呈现主要变量间的关系：

                                        | 关系类型 | 源变量 | 目标变量 | 统计量 | p值 | 显著性 | 关系强度 |
                                        | --- | --- | --- | --- | --- | --- | --- |

                                        ## 详细关系分析
                                        针对每个重要关系，使用以下格式：

                                        ### 关系: [源变量] -> [目标变量]
                                        **统计结果:**
                                        - 关系类型: [分类->数值/数值->数值/分类->分类]
                                        - 统计量: [F值/相关系数/卡方值] = X
                                        - p值: X
                                        - 显著性: [高度显著/显著/不显著]

                                        **类别统计:** (适用于分类->数值关系)
                                        | 类别 | 样本数 | 均值 | 标准差 | 与总体均值差异 |
                                        | --- | --- | --- | --- | --- |

                                        **显著差异类别:**
                                        - 类别X: 均值Y (比总体均值[高/低]Z%)

                                        # 技术附录
                                        ## DAG分析完整报告
                                        [此处直接插入原始DAG报告，不做修改]

                                        请确保：
                                        1. 表格格式清晰，数据对齐
                                        2. 所有统计结果保留适当小数位数(通常2-4位)
                                        3. 对统计显著性使用标准表示: *** p<0.001, ** p<0.01, * p<0.05, ns p≥0.05
                                        4. 对非技术读者解释统计术语，但保持专业性
                                        5. 重点突出异常值、显著差异和有商业价值的发现
                                        6. 所有结论必须有数据支持，避免过度解读
                                        7. 对于复杂关系，提供简明的解释和可能的因果机制
                                        8. 请提供尽可能详细的数据洞察和业务影响分析"""
                                },
                                {
                                    "role": "user",
                                    "content": f"""
                                    初始分析结果：{st.session_state.analysis_response}
                                    DAG分析报告：{dag_report}
                                    数据描述信息：{json_output}
                                    
                                    请基于以上信息生成一份专业的数据分析报告。报告应严格遵循系统提示中的结构和格式要求，特别注意：
                                    1. 充分利用DAG分析报告中的统计结果，包括F值、p值和类别统计
                                    2. 整合数据描述中的统计信息和分布特征
                                    3. 所有表格必须格式规范，数据对齐
                                    4. 重点关注统计显著的关系和异常值
                                    5. 确保非技术人员也能理解报告内容
                                    6. 所有结论和建议必须有数据支持
                                    """
                                }
                            ],
                            temperature=0.7,
                            max_tokens=5000,
                            stream=True
                        )
                        
                        full_response = ""
                        reasoning_content = ""
                        
                        # 处理流式响应
                        for chunk in response:
                            if hasattr(chunk.choices[0].delta, 'reasoning_content') and chunk.choices[0].delta.reasoning_content:
                                reasoning_content += chunk.choices[0].delta.reasoning_content
                                dag_progress.markdown(f"思考过程：\n{reasoning_content}\n\n分析结果：\n{full_response}")
                            elif hasattr(chunk.choices[0].delta, 'content') and chunk.choices[0].delta.content:
                                full_response += chunk.choices[0].delta.content
                                dag_progress.markdown(f"思考过程：\n{reasoning_content}\n\n分析结果：\n{full_response}")
                        
                        # 保存到session state
                         # --- 新增的清理步骤 ---
                        cleaned_report = full_response.strip()
                        # 检查并移除包裹的代码块标记
                        if cleaned_report.startswith("```markdown"):
                            cleaned_report = cleaned_report[len("```markdown"):].strip()
                        if cleaned_report.startswith("```"):
                            cleaned_report = cleaned_report[3:].strip()
                        if cleaned_report.endswith("```"):
                            cleaned_report = cleaned_report[:-3].strip()

                        # 保存清理后的报告到session state
                        st.session_state.business_report = cleaned_report  # <-- 使用清理后的变量
                
                        # st.session_state.business_report = full_response
                        st.session_state.dag_report = dag_report  # 确保这行存在，你的代码里已经有了
                        st.session_state.data_description = json_output # 确保这行存在，你的代码里已经有了
                        st.session_state.dag_reasoning = reasoning_content
                        
                        # 清空进度容器
                        dag_progress.empty()
                        
                except Exception as e:
                    st.error(f"DAG分析过程中出错: {str(e)}")
        
        # 使用两列显示分析结果和DAG分析结果
        if st.session_state.analysis_response or st.session_state.business_report:
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("### 数据分析")
                if st.session_state.analysis_reasoning:
                    with st.expander("思考过程", expanded=False):
                        st.markdown(st.session_state.analysis_reasoning)
                st.markdown(st.session_state.analysis_response)
                
                # DAG定义编辑器
                st.markdown("### DAG定义")
                dag_editor = st.text_area("编辑DAG边定义:", value=st.session_state.dag_edges, height=200)
                if dag_editor != st.session_state.dag_edges:
                    st.session_state.dag_edges = dag_editor
                # Generate the Mermaid HTML
                try:
                    # 确保先解析文本中的dag_edges
                    dag_edges = extract_dag_edges(st.session_state.dag_edges)
                    if dag_edges:
                        mermaid_html = create_mermaid_html_from_edges(dag_edges)
                        # 渲染 Mermaid 图，并增加高度以适应缩放
                        st.markdown("## 关系图 (使用右上角按钮缩放)")
                        st.components.v1.html(mermaid_html, height=600, scrolling=True)
                    else:
                        st.warning("无法解析DAG边定义，请检查格式是否正确")
                except Exception as e:
                    st.error(f"生成关系图时出错: {str(e)}")
            
            with col2:
                if st.session_state.business_report:
                    st.markdown("### DAG分析结果")

                    # 显示LLM的分析过程
                    if hasattr(st.session_state, 'dag_reasoning') and st.session_state.dag_reasoning:
                        with st.expander("LLM分析过程", expanded=False):
                            st.markdown(st.session_state.dag_reasoning)
                    
                    # 1. 使用 st.markdown() 显示格式化的商业报告
                    st.markdown(st.session_state.business_report)

                    # 2. 使用 st.expander + st.text() 显示原始的DAG技术报告
                    if hasattr(st.session_state, 'dag_report') and st.session_state.dag_report:
                        with st.expander("查看DAG分析原始报告 (Technical Appendix)", expanded=False):
                            # st.text() 能很好地保留原始文本的对齐和换行
                            st.text(st.session_state.dag_report)

                    # 3. 使用 st.expander + st.json() 显示描述性统计的JSON数据
                    if hasattr(st.session_state, 'data_description') and st.session_state.data_description:
                        with st.expander("查看描述性统计JSON数据", expanded=False):
                            # st.json() 会美化显示JSON
                            st.json(st.session_state.data_description)
            
            # 添加下载按钮
            if st.session_state.business_report:
                st.markdown("---")
                st.download_button(
                    label="📥 下载分析报告",
                    data=create_word_document(st.session_state.business_report),
                    file_name=f"data_analysis_report_{time.strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

# Add this function before the main() function
def setup_sales_forecasting():
    st.markdown(
        """
    <h1 style='text-align: center; font-size: 18px; font-weight: bold;'>Sales Forecasting Application</h1>
    <h6 style='text-align: center; font-size: 12px;'>Upload your Excel file to forecast sales or other time series data</h6>
    <br><br><br>
    """,
        unsafe_allow_html=True,
    )

    # Set random seed for reproducibility
    np.random.seed(42)
        
    # Initialize session state for storing results between interactions
    if 'forecast_df' not in st.session_state:
        st.session_state.forecast_df = None
    if 'all_groups' not in st.session_state:
        st.session_state.all_groups = []
    if 'has_forecast' not in st.session_state:
        st.session_state.has_forecast = False
    if 'target_column' not in st.session_state:
        st.session_state.target_column = None
    if 'original_filename' not in st.session_state:
        st.session_state.original_filename = None
    if 'training_accuracy' not in st.session_state:
        st.session_state.training_accuracy = {}
    if 'validation_accuracy' not in st.session_state:
        st.session_state.validation_accuracy = {}
    if 'start_date_str' not in st.session_state:
        st.session_state.start_date_str = None
    if 'training_end_date_str' not in st.session_state:
        st.session_state.training_end_date_str = None
    if 'covariate_columns' not in st.session_state:
        st.session_state.covariate_columns = []
    
    st.markdown("""
    This app helps you forecast sales or other time series data using Facebook Prophet.
    Upload your Excel file, configure the settings, and get predictions for future periods.
    """)
    
    # File uploader
    uploaded_file = st.file_uploader("Upload Excel file", type=["xlsx", "xls"])
    
    if uploaded_file is not None:
        # Store original filename
        if st.session_state.original_filename is None:
            st.session_state.original_filename = uploaded_file.name.split('.')[0]
            
        # Load the file and display sheet selection
        excel_file = pd.ExcelFile(uploaded_file)
        sheet_name = st.selectbox("Select Sheet", excel_file.sheet_names)
        
        # Read the selected sheet
        df = pd.read_excel(uploaded_file, sheet_name=sheet_name)
        
        # Display data preview
        st.subheader("Data Preview")
        st.dataframe(df.head())
        
        # Basic data info
        st.subheader("Dataset Information")
        col1, col2 = st.columns(2)
        with col1:
            st.write(f"Number of rows: {df.shape[0]}")
        with col2:
            st.write(f"Number of columns: {df.shape[1]}")
        
        # Column selection and configuration
        st.subheader("Configure Forecasting Parameters")
        
        col1, col2 = st.columns(2)
        
        with col1:
            # Date column selection
            date_column_options = df.columns.tolist()
            date_column = st.selectbox("Select Date/Month Column", date_column_options)
            
            # Date format selection
            date_format_options = [
                "%Y%m", "%Y-%m", "%b %Y", "%B %Y", 
                "%Y%m%d", "%Y-%m-%d", "%d-%m-%Y", "%m/%d/%Y"
            ]
            date_format = st.selectbox("Select Date Format", date_format_options)
            
            # 确定预测频率 - 根据日期格式自动调整
            forecast_freq = 'D' if any(x in date_format for x in ['%d', '%Y%m%d', '%Y-%m-%d', '%d-%m-%Y', '%m/%d/%Y']) else 'MS'
            
            # Target column (to be predicted)
            numeric_columns = df.select_dtypes(include=['number']).columns.tolist()
            target_column = st.selectbox("Select Target Column to Forecast", numeric_columns)
            st.session_state.target_column = target_column
    
            # Start date for training
            min_date = pd.to_datetime("2020-01-01")  # Default minimum date
            max_date = datetime.now()
            start_date = st.date_input("Training Start Date", 
                                      value=pd.to_datetime("2022-01-01"),
                                      min_value=min_date,
                                      max_value=max_date)
            
            # Covariate selection - NEW FEATURE
            # Filter out date column and target column from potential covariates
            potential_covariates = [col for col in numeric_columns if col != target_column]
            covariate_columns = st.multiselect("Select Covariate Columns (Optional)", potential_covariates, 
                                               help="Additional numeric variables to improve forecast accuracy")
            st.session_state.covariate_columns = covariate_columns
            
            # 添加模型选择
            use_advanced_model = st.checkbox("使用高级模型 (更高精度但更慢)", value=False)   
        
        with col2:
            # Grouping columns (multi-select)
            all_columns = df.columns.tolist()
            grouping_columns = st.multiselect("Select Columns for Grouping (Optional)", all_columns)
    
            # Forecast horizon
            forecast_years = st.number_input("Forecast Horizon (Years)", min_value=1, max_value=10, value=3)
            
            # Add a forecast end date input
            current_date = datetime.now()
            default_end_date = current_date.replace(day=1) + pd.DateOffset(years=forecast_years) - pd.DateOffset(days=1)
            end_date = st.date_input("Forecast End Date", 
                                    value=default_end_date,
                                    min_value=current_date)
            
            # Hidden Confidence Interval Width (not displayed in UI)
            interval_width = 0.6  # Set default to 0.6 as requested
            
            # Add a training end date input at the bottom of col2
            training_end_date = st.date_input("Training End Date", 
                                    value=current_date,
                                    min_value=start_date,
                                    max_value=current_date)
        
        # Run forecast button
        if st.button("Run Forecast") or st.session_state.has_forecast:
            if not st.session_state.has_forecast:
                try:
                    # Progress bar
                    progress_bar = st.progress(0)
                    
                    # Copy the dataframe to avoid modifying the original
                    df_copy = df.copy()
                    
                    # Process date column
                    st.info("Processing date column...")
                    try:
                        df_copy['date'] = pd.to_datetime(df_copy[date_column].astype(str), format=date_format)
                    except:
                        df_copy['date'] = pd.to_datetime(df_copy[date_column].astype(str))
                        
                    # Store the original date format from the column
                    original_date_values = df_copy[date_column].copy()
                        
                    progress_bar.progress(10)
                    
                    # Create grouping key if group columns are selected
                    if grouping_columns:
                        st.info("Creating group combinations...")
                        df_copy['group'] = df_copy[grouping_columns[0]].astype(str)
                        for col in grouping_columns[1:]:
                            df_copy['group'] += '_' + df_copy[col].astype(str)
                    else:
                        # If no grouping is selected, create a single group
                        df_copy['group'] = 'all_data'
                        
                    progress_bar.progress(20)
                    
                    # Get current date and calculate end date for forecasting
                    current_date = datetime.now()
                    end_date_str = end_date.strftime('%Y-%m-%d')
                    training_end_date_str = training_end_date.strftime('%Y-%m-%d')
                    start_date_str = start_date.strftime('%Y-%m-%d')
                    
                    # Store date strings in session state for later use
                    st.session_state.start_date_str = start_date_str
                    st.session_state.training_end_date_str = training_end_date_str
                    
                    # Create complete date and group combinations
                    # 创建完整日期和分组组合，使用动态频率
                    st.info("Creating complete date-group combinations...")
                    # 确保生成的日期范围始终延伸到预测结束日期
                    all_dates = pd.date_range(start=df_copy['date'].min(), end=end_date_str, freq=forecast_freq)
                    all_groups = df_copy['group'].unique()
                    st.session_state.all_groups = all_groups.tolist()
                    
                    complete_df = pd.DataFrame([(date, group) for date in all_dates for group in all_groups],
                                              columns=['date', 'group'])
                    
                    # Merge with original data
                    df_copy = pd.merge(complete_df, df_copy, on=['date', 'group'], how='left')
                    
                    # Fill missing values for target column with 0
                    df_copy[target_column] = df_copy[target_column].fillna(0)
                    
                    # Handle covariates - calculate means for each covariate within its group
                    covariate_means = {}
                    if covariate_columns:
                        for col in covariate_columns:
                            # Calculate mean for each group
                            for group in all_groups:
                                group_data = df_copy[(df_copy['group'] == group) & df_copy[col].notna()]
                                if not group_data.empty:
                                    # Use the mean of non-NA values for this group
                                    covariate_means[(group, col)] = group_data[col].mean()
                                else:
                                    # If all values are NA for this group, use global mean
                                    global_mean = df_copy[df_copy[col].notna()][col].mean()
                                    covariate_means[(group, col)] = global_mean if not np.isnan(global_mean) else 0
                            
                            # Fill NAs with group-specific means
                            for group in all_groups:
                                mask = (df_copy['group'] == group) & df_copy[col].isna()
                                df_copy.loc[mask, col] = covariate_means.get((group, col), 0)
                    
                    # Forward fill other columns within groups
                    if grouping_columns:
                        for col in grouping_columns:
                            df_copy[col] = df_copy.groupby('group')[col].ffill().bfill()
                            
                    progress_bar.progress(40)
                    
                    # Sort data
                    df_copy = df_copy.sort_values(by=['group', 'date'], ascending=[True, True])
                    
                    # Convert dates to string format
                    current_date_str = current_date.strftime("%Y-%m-%d")
                    
                    # Run Prophet forecast for each group
                    st.info("Running forecasts for each group...")
                    forecasts = {}
                    training_accuracy = {}
                    validation_accuracy = {}
                    total_groups = len(all_groups)
                    
                    for i, group in enumerate(all_groups):
                        # Update progress based on group progress
                        progress_value = 40 + (i / total_groups * 50)
                        progress_bar.progress(int(progress_value))
                        
                        # Filter data for this group - using the start_date parameter
                        group_data = df_copy[(df_copy['group'] == group) & 
                                           (df_copy['date'] >= start_date_str) & 
                                           (df_copy['date'] <= training_end_date_str)].copy()
                        
                        group_data = group_data.rename(columns={'date': 'ds', target_column: 'y'})
                        
                        # Add covariates to Prophet model if selected
                        if use_advanced_model:
                            model = Prophet(interval_width=interval_width, uncertainty_samples=1000, mcmc_samples=300)
                        else:
                            model = Prophet(interval_width=interval_width)
                            
                        # Add covariates to model
                        if covariate_columns:
                            for col in covariate_columns:
                                model.add_regressor(col)
    
                        try:
                            model.fit(group_data[['ds', 'y'] + covariate_columns])
                            
                            # 修改: 直接使用预测结束日期来创建future dataframe
                            future_end_date = pd.to_datetime(end_date_str)
                            # 创建从训练数据开始到预测结束日期的完整日期范围，使用动态频率
                            future_dates = pd.date_range(start=group_data['ds'].min(), end=future_end_date, freq=forecast_freq)
                            future = pd.DataFrame({'ds': future_dates})
                            
                            # Add covariate values to future dataframe for prediction
                            if covariate_columns:
                                for col in covariate_columns:
                                    # Get all values for this group and covariate
                                    historical_values = df_copy[(df_copy['group'] == group) & 
                                                              (df_copy['date'].isin(future_dates))][col].values
                                    
                                    # Create DataFrame with future dates and corresponding covariate values
                                    covariate_df = pd.DataFrame({
                                        'ds': future_dates[:len(historical_values)],
                                        col: historical_values
                                    })
                                    
                                    # For dates beyond historical data, use the group mean
                                    if len(future_dates) > len(historical_values):
                                        remaining_dates = future_dates[len(historical_values):]
                                        remaining_values = np.full(len(remaining_dates), 
                                                                  covariate_means.get((group, col), 0))
                                        
                                        remaining_df = pd.DataFrame({
                                            'ds': remaining_dates,
                                            col: remaining_values
                                        })
                                        
                                        covariate_df = pd.concat([covariate_df, remaining_df])
                                    
                                    # Merge with future DataFrame
                                    future = pd.merge(future, covariate_df, on='ds', how='left')
                            
                            # Make forecast
                            forecast = model.predict(future)
                            
                            forecasts[group] = forecast[['ds', 'yhat', 'yhat_lower', 'yhat_upper']]
                        except Exception as e:
                            st.warning(f"Could not forecast for group {group}: {str(e)}")
                            # Create empty forecast data for this group
                            forecasts[group] = pd.DataFrame({
                                'ds': all_dates,
                                'yhat': np.zeros(len(all_dates)),
                                'yhat_lower': np.zeros(len(all_dates)),
                                'yhat_upper': np.zeros(len(all_dates))
                            })
                    
                    progress_bar.progress(90)
                    
                    # Write forecasts back to the dataframe using direct array assignment
                    st.info("Combining results...")
                    for group, forecast in forecasts.items():
                        # Use date matching approach instead of direct array assignment
                        for i, row in forecast.iterrows():
                            forecast_date = row['ds']
                            # Match both group and date
                            mask = (df_copy['group'] == group) & (df_copy['date'] == forecast_date)
                            # Assign forecast values
                            df_copy.loc[mask, 'forecast'] = row['yhat']
                            df_copy.loc[mask, 'forecast_lower'] = row['yhat_lower']
                            df_copy.loc[mask, 'forecast_upper'] = row['yhat_upper']
                    
                    # Format dates in the original date column according to the selected format
                    # This preserves the original date format while extending to future dates
                    df_copy[date_column] = df_copy['date'].dt.strftime(date_format)
                    
                    # Calculate prediction accuracy percentages for each group
                    for group in all_groups:
                        # Training period accuracy (Training Start Date to Training End Date)
                        training_data = df_copy[(df_copy['group'] == group) & 
                                              (df_copy['date'] >= start_date_str) & 
                                              (df_copy['date'] <= training_end_date_str)]
                        
                        if not training_data.empty and training_data[target_column].sum() > 0:
                            actual = training_data[target_column].values
                            pred = training_data['forecast'].values
                            # Calculate percentage error: (pred - actual) / actual
                            pct_errors = np.where(actual > 0, (pred - actual) / actual, np.nan)
                            # Calculate MAPE (Mean Absolute Percentage Error)
                            mean_pct_error = np.nanmean(pct_errors) * 100
                            training_accuracy[group] = mean_pct_error
                        else:
                            training_accuracy[group] = np.nan
                        
                        # Validation period accuracy (Training End Date to current date)
                        validation_data = df_copy[(df_copy['group'] == group) & 
                                                (df_copy['date'] > training_end_date_str) & 
                                                (df_copy['date'] <= current_date_str)]
                        
                        if not validation_data.empty and validation_data[target_column].sum() > 0:
                            actual = validation_data[target_column].values
                            pred = validation_data['forecast'].values
                            # Calculate percentage error: (pred - actual) / actual
                            pct_errors = np.where(actual > 0, (pred - actual) / actual, np.nan)
                            # Calculate MAPE
                            mean_pct_error = np.nanmean(pct_errors) * 100
                            validation_accuracy[group] = mean_pct_error
                        else:
                            validation_accuracy[group] = np.nan
                    
                    # Store accuracy metrics in session state
                    st.session_state.training_accuracy = training_accuracy
                    st.session_state.validation_accuracy = validation_accuracy
                    
                    # Store results in session state
                    st.session_state.forecast_df = df_copy
                    st.session_state.has_forecast = True
                    
                    progress_bar.progress(100)
                    st.success("Forecast completed successfully!")
                    
                except Exception as e:
                    st.error(f"An error occurred during forecasting: {str(e)}")
                    st.error("Please check your data and settings and try again.")
            
            # Display results if available in session state
            if st.session_state.has_forecast and st.session_state.forecast_df is not None:
                # Display results
                st.subheader("Forecast Results")
                st.dataframe(st.session_state.forecast_df)
                
                # Create download link for the forecast
                st.subheader("Download Complete Forecast")
                
                # Get the original filename for downloads
                filename_base = st.session_state.original_filename
                
                # CSV download with custom filename
                csv = st.session_state.forecast_df.to_csv(index=False)
                b64 = base64.b64encode(csv.encode()).decode()
                csv_filename = f"{filename_base}_forecast.csv"
                href = f'<a href="data:file/csv;base64,{b64}" download="{csv_filename}">Download CSV File</a>'
                st.markdown(href, unsafe_allow_html=True)
                
                # Excel download with custom filename
                buffer = io.BytesIO()
                st.session_state.forecast_df.to_excel(buffer, index=False)
                buffer.seek(0)
                b64_excel = base64.b64encode(buffer.read()).decode()
                excel_filename = f"{filename_base}_forecast.xlsx"
                href_excel = f'<a href="data:application/vnd.openxmlformats-officedocument.spreadsheetml.sheet;base64,{b64_excel}" download="{excel_filename}">Download Excel File</a>'
                st.markdown(href_excel, unsafe_allow_html=True)
                
                # Visualization section
                st.subheader("Forecast Visualizations")
                
                # Create dropdown for group selection
                selected_group = st.selectbox("Select group to visualize", st.session_state.all_groups)
                
                # Get accuracy metrics for selected group
                training_acc = st.session_state.training_accuracy.get(selected_group, np.nan)
                validation_acc = st.session_state.validation_accuracy.get(selected_group, np.nan)
                
                # Format accuracy metrics for display
                training_acc_text = f"Training Accuracy: {training_acc:.2f}%" if not np.isnan(training_acc) else "Training Accuracy: N/A"
                validation_acc_text = f"Validation Accuracy: {validation_acc:.2f}%" if not np.isnan(validation_acc) else "Validation Accuracy: N/A"
                
                # Plot the selected group
                selected_data = st.session_state.forecast_df[st.session_state.forecast_df['group'] == selected_group].copy()
                
                # 修改: 确保使用排序后的数据来创建图表，避免confidence interval的显示问题
                selected_data = selected_data.sort_values('date')
                
                # Create plotly figure for selected group
                fig = make_subplots(specs=[[{"secondary_y": True}]])
                
                # Add actual values
                fig.add_trace(
                    go.Scatter(
                        x=selected_data['date'],
                        y=selected_data[st.session_state.target_column],
                        mode='lines+markers',
                        name='Actual',
                        line=dict(color='blue')
                    )
                )
                
                # Add forecast
                fig.add_trace(
                    go.Scatter(
                        x=selected_data['date'],
                        y=selected_data['forecast'],
                        mode='lines',
                        name='Forecast',
                        line=dict(color='red')
                    )
                )
                
                # 修改: 更改confidence interval的绘制方法，确保正确绘制
                # 添加上边界
                fig.add_trace(
                    go.Scatter(
                        x=selected_data['date'],
                        y=selected_data['forecast_upper'],
                        mode='lines',
                        line=dict(width=0),
                        showlegend=False
                    )
                )
                
                # 添加下边界
                fig.add_trace(
                    go.Scatter(
                        x=selected_data['date'],
                        y=selected_data['forecast_lower'],
                        mode='lines',
                        line=dict(width=0),
                        fillcolor='rgba(255,0,0,0.2)',
                        fill='tonexty',
                        name='Confidence Interval'
                    )
                )
                
                # Update layout with accuracy metrics in the title
                fig.update_layout(
                    title=f'Forecast for {selected_group} ({training_acc_text}, {validation_acc_text})',
                    xaxis_title='Date',
                    yaxis_title=st.session_state.target_column,
                    hovermode='x unified',
                    legend=dict(orientation='h', y=1.1)
                )
                
                st.plotly_chart(fig, use_container_width=True)
                
                # Add option to view overall forecast (sum of all groups)
                if st.checkbox("View Overall Forecast (Sum of All Groups)"):
                    # Group by date and sum the values
                    overall_data = st.session_state.forecast_df.groupby('date').agg({
                        st.session_state.target_column: 'sum',
                        'forecast': 'sum',
                        'forecast_lower': 'sum',
                        'forecast_upper': 'sum'
                    }).reset_index()
                    
                    # 修改: 确保使用排序后的数据来创建图表，避免confidence interval的显示问题
                    overall_data = overall_data.sort_values('date')
                    
                    # Calculate overall accuracy metrics
                    # Convert dates to datetime for comparison
                    overall_data['date_dt'] = pd.to_datetime(overall_data['date'])
                    
                    # Use stored date strings from session state
                    start_date_str = st.session_state.start_date_str
                    training_end_date_str = st.session_state.training_end_date_str
                    
                    # Convert to datetime objects for comparison
                    start_date_dt = pd.to_datetime(start_date_str)
                    training_end_date_dt = pd.to_datetime(training_end_date_str)
                    current_date_dt = pd.to_datetime(current_date.strftime("%Y-%m-%d"))
                    
                    # Training period accuracy (Training Start Date to Training End Date)
                    training_overall = overall_data[(overall_data['date_dt'] >= start_date_dt) & 
                                                  (overall_data['date_dt'] <= training_end_date_dt)]
                    
                    if not training_overall.empty and training_overall[st.session_state.target_column].sum() > 0:
                        actual = training_overall[st.session_state.target_column].values
                        pred = training_overall['forecast'].values
                        # Calculate percentage error: (pred - actual) / actual
                        pct_errors = np.where(actual > 0, (pred - actual) / actual, np.nan)
                        # Calculate MAPE
                        overall_training_acc = np.nanmean(pct_errors) * 100
                    else:
                        overall_training_acc = np.nan
                    
                    # Validation period accuracy (Training End Date to current date)
                    validation_overall = overall_data[(overall_data['date_dt'] > training_end_date_dt) & 
                                                    (overall_data['date_dt'] <= current_date_dt)]
                    
                    if not validation_overall.empty and validation_overall[st.session_state.target_column].sum() > 0:
                        actual = validation_overall[st.session_state.target_column].values
                        pred = validation_overall['forecast'].values
                        # Calculate percentage error: (pred - actual) / actual
                        pct_errors = np.where(actual > 0, (pred - actual) / actual, np.nan)
                        # Calculate MAPE
                        overall_validation_acc = np.nanmean(pct_errors) * 100
                    else:
                        overall_validation_acc = np.nan
                    
                    # Format overall accuracy metrics for display
                    overall_training_acc_text = f"Training Accuracy: {overall_training_acc:.2f}%" if not np.isnan(overall_training_acc) else "Training Accuracy: N/A"
                    overall_validation_acc_text = f"Validation Accuracy: {overall_validation_acc:.2f}%" if not np.isnan(overall_validation_acc) else "Validation Accuracy: N/A"
                    
                    # Create plotly figure for overall data
                    fig_overall = make_subplots(specs=[[{"secondary_y": True}]])
                    
                    # Add actual values
                    fig_overall.add_trace(
                        go.Scatter(
                            x=overall_data['date'],
                            y=overall_data[st.session_state.target_column],
                            mode='lines+markers',
                            name='Actual',
                            line=dict(color='blue')
                        )
                    )
                    
                    # Add forecast
                    fig_overall.add_trace(
                        go.Scatter(
                            x=overall_data['date'],
                            y=overall_data['forecast'],
                            mode='lines',
                            name='Forecast',
                            line=dict(color='red')
                        )
                    )
                    
                    # 修改: 更改confidence interval的绘制方法，确保正确绘制
                    # 添加上边界
                    fig_overall.add_trace(
                        go.Scatter(
                            x=overall_data['date'],
                            y=overall_data['forecast_upper'],
                            mode='lines',
                            line=dict(width=0),
                            showlegend=False
                        )
                    )
                    
                    # 添加下边界
                    fig_overall.add_trace(
                        go.Scatter(
                            x=overall_data['date'],
                            y=overall_data['forecast_lower'],
                            mode='lines',
                            line=dict(width=0),
                            fillcolor='rgba(255,0,0,0.2)',
                            fill='tonexty',
                            name='Confidence Interval'
                        )
                    )
                    
                    # Update layout with accuracy metrics in the title
                    fig_overall.update_layout(
                        title=f'Overall Forecast (Sum of All Groups) ({overall_training_acc_text}, {overall_validation_acc_text})',
                        xaxis_title='Date',
                        yaxis_title=st.session_state.target_column,
                        hovermode='x unified',
                        legend=dict(orientation='h', y=1.1)
                    )
                    
                    st.plotly_chart(fig_overall, use_container_width=True)
        
        # Add button to clear results and start over
        if st.session_state.has_forecast:
            if st.button("Clear Results and Start Over"):
                # Fix for experimental_rerun issue - use session state to clear data
                st.session_state.forecast_df = None
                st.session_state.all_groups = []
                st.session_state.has_forecast = False
                st.session_state.target_column = None
                st.session_state.original_filename = None
                st.session_state.training_accuracy = {}
                st.session_state.validation_accuracy = {}
                st.session_state.start_date_str = None
                st.session_state.training_end_date_str = None
                # Use Streamlit's rerun function instead of experimental_rerun
                st.rerun()
    
    else:
        # Display sample instructions when no file is uploaded
        st.info("Please upload an Excel file to begin. Your file should contain:")
        st.markdown("""
        - A column with dates or year-month values
        - At least one numeric column to forecast
        - Optional grouping columns (e.g., product codes, regions)
        
        The app will:
        1. Process your data
        2. Fill missing values
        3. Generate forecasts using Facebook Prophet
        4. Allow you to download the complete forecast results
        """)
        
        # Sample data structure
        st.subheader("Sample Data Structure")
        sample_data = pd.DataFrame({
            'year_month': ['202201', '202202', '202203', '202201', '202202', '202203'],
            'product_code': ['A001', 'A001', 'A001', 'B002', 'B002', 'B002'],
            'region': ['North', 'North', 'North', 'South', 'South', 'South'],
            'sales_quantity': [100, 120, 110, 80, 85, 95],
            'sales_amount': [5000, 6000, 5500, 4000, 4250, 4750]
        })
        
        st.dataframe(sample_data)
    
    # Footer
    st.markdown("""
    ---
    ### Notes:
    - For best results, provide at least 12 months of historical data
    - If using grouping, ensure all groups have sufficient data points
    - The forecast horizon can be adjusted based on your needs
    """)

# 在文件顶部的导入部分下添加这个辅助函数
def convert_cluster_label(label, naming_style):
    """将数字标签转换为选定的命名风格"""
    label_num = int(label)
    if naming_style == "数字":
        return str(label_num)
    elif naming_style == "字母":
        # 将 0, 1, 2, ... 转换为 A, B, C, ...
        return chr(65 + label_num) if label_num < 26 else f"A{label_num-25}"
    elif naming_style == "中文":
        # 将 0, 1, 2, ... 转换为 甲, 乙, 丙, ...
        chinese_labels = ['甲', '乙', '丙', '丁', '戊', '己', '庚', '辛', '壬', '癸']
        if label_num < len(chinese_labels):
            return chinese_labels[label_num]
        else:
            # 超出基本的10个后使用"甲1, 甲2, ..."格式
            return f"{chinese_labels[0]}{label_num-9}"
    return str(label_num)  # 默认返回数字

def setup_clustering_analysis():
    st.markdown(
        """
    <h1 style='text-align: center; font-size: 18px; font-weight: bold;'>K-means/DBSCAN聚类与热力图分析工具</h1>
    <h6 style='text-align: center; font-size: 12px;'>上传CSV或Excel文件进行K-means/DBSCAN聚类与热力图分析</h6>
    <br><br><br>
    """,
        unsafe_allow_html=True,
    )

    # 应用标题
    st.title("K-means/DBSCAN聚类与热力图分析工具")

    # 边栏说明
    with st.sidebar:
        st.header("使用说明")
        st.write("""
        1. 上传 CSV 或 Excel 数据文件
        2. 选择聚类方式:
        - 使用 K-means 或 DBSCAN 进行聚类
        - 或直接使用已有的聚类列
        3. 观察热力图分析结果
        4. 下载分析结果
        """)

    # 上传文件
    uploaded_file = st.file_uploader("上传数据文件", type=["csv", "xlsx", "xls"])

    if uploaded_file is not None:
        # 读取数据
        try:
            if uploaded_file.name.endswith('.csv'):
                df = pd.read_csv(uploaded_file)
            else:
                df = pd.read_excel(uploaded_file)
            
            st.success("文件上传成功!")
            
            # 显示数据预览
            st.subheader("数据预览")
            st.dataframe(df.head())
            
            # 聚类模式选择
            clustering_mode = st.radio(
                "选择聚类模式:",
                ["使用模型聚类", "使用已有聚类列"]
            )
            
            # 用于存储两种聚类结果的变量
            cluster_col1 = None
            cluster_col2 = None
            
            # 第一个聚类
            st.subheader("第一个聚类")
            
            if clustering_mode == "使用模型聚类":
                # 选择聚类算法
                clustering_algo = st.selectbox("选择聚类算法", ["K-means", "DBSCAN"])
                
                # 添加簇命名风格选择
                naming_style1 = st.selectbox(
                    "选择簇命名风格",
                    ["字母", "数字", "中文"],
                    index=0,  # 默认选择字母
                    key="naming_style1"
                )
                # 选择用于聚类的列
                numeric_cols = df.select_dtypes(include=['int64', 'float64']).columns.tolist()
                if len(numeric_cols) > 0:
                    selected_cols1 = st.multiselect(
                        "选择用于第一个聚类的列",
                        numeric_cols,
                        default=numeric_cols[:max(0,min(3, len(numeric_cols)))]
                    )
                    
                    if selected_cols1:
                        X1 = df[selected_cols1]
                        # 标准化数据
                        scaler = StandardScaler()
                        X_scaled1 = scaler.fit_transform(X1)
                        
                        # 多维数据降维选项
                        if len(selected_cols1) >= 2:
                            input_option1 = st.radio(
                                "选择聚类输入数据（第一个聚类）",
                                ["原始数据", "t-SNE降维到2维"],
                                index=0
                            )
                            if input_option1 == "t-SNE降维到2维":
                                tsne = TSNE(n_components=2, random_state=42)
                                X_scaled1 = tsne.fit_transform(X_scaled1)
                        
                        if clustering_algo == "K-means":
                            # 设置 K 值
                            k_value1 = st.slider("选择 K-means 的 K 值（簇的数量）", 2, 10, 3)
                            
                            # 执行 K-means
                            kmeans1 = KMeans(n_clusters=k_value1, random_state=42, n_init=10)
                            # 替换为:
                            labels = kmeans1.fit_predict(X_scaled1)
                            df['Cluster1'] = [convert_cluster_label(str(label), naming_style1) for label in labels]                        
                            # 计算轮廓系数
                            silhouette_avg = silhouette_score(X_scaled1, df['Cluster1'])
                            st.write(f"轮廓系数 (越接近1越好): {silhouette_avg:.4f}")
                        
                        else:  # DBSCAN
                            # 设置 DBSCAN 参数并添加通俗说明
                            st.markdown("""
                            **DBSCAN 参数说明**:
                            - **eps**: 点的邻域半径，决定多远的点被认为是“邻居”。值越小，簇越密集；值越大，簇越分散。
                            - **min_samples**: 一个簇所需的最小点数（包括核心点本身）。值越大，簇需要更多点才能形成；值越小，可能生成更多小簇。
                            """)
                            eps = st.slider("选择 DBSCAN 的 eps 参数", 0.1, 2.0, 0.5, step=0.01)
                            min_samples = st.slider("选择 DBSCAN 的 min_samples 参数", 2, 20, 5)
                            
                            # 执行 DBSCAN
                            dbscan1 = DBSCAN(eps=eps, min_samples=min_samples)
                            labels = dbscan1.fit_predict(X_scaled1)
                            
                            # 处理噪声点 (-1)
                            if -1 in labels:
                                # 找到非噪声点
                                non_noise_mask = labels != -1
                                X_non_noise = X_scaled1[non_noise_mask]
                                labels_non_noise = labels[non_noise_mask]
                                
                                # 对噪声点使用 KNN 分配最近的簇
                                noise_mask = labels == -1
                                X_noise = X_scaled1[noise_mask]
                                
                                if len(X_non_noise) > 0 and len(X_noise) > 0:
                                    knn = NearestNeighbors(n_neighbors=1)
                                    knn.fit(X_non_noise)
                                    _, indices = knn.kneighbors(X_noise)
                                    labels[noise_mask] = labels_non_noise[indices.flatten()]
                            
                            df['Cluster1'] = [convert_cluster_label(str(label), naming_style1) for label in labels]

                            # 检查是否有有效簇
                            if len(np.unique(df['Cluster1'])) > 1:
                                silhouette_avg = silhouette_score(X_scaled1, df['Cluster1'])
                                st.write(f"轮廓系数 (越接近1越好): {silhouette_avg:.4f}")
                            else:
                                st.warning("DBSCAN 未能生成多个有效簇，请调整 eps 或 min_samples 参数")
                        
                        # 显示聚类结果
                        st.write("聚类结果分布:")
                        cluster_counts1 = df['Cluster1'].value_counts().sort_index()
                        st.write(cluster_counts1)
                        
                        # 可视化聚类结果
                        if len(selected_cols1) == 1:
                            # 一维：分布图 + 簇分割线
                            fig, ax = plt.subplots(figsize=(10, 6))
                            unique_clusters = sorted(df['Cluster1'].unique())
                            colors = sns.color_palette("husl", len(unique_clusters))
                            
                            for cluster, color in zip(unique_clusters, colors):
                                cluster_data = X1[selected_cols1[0]][df['Cluster1'] == cluster]
                                sns.histplot(cluster_data, kde=True, label=f'簇 {cluster}', 
                                        stat='density', alpha=0.4, color=color, ax=ax)
                            
                            # 添加簇分割线（基于簇中心）
                            cluster_centers = []
                            for cluster in unique_clusters:
                                cluster_data = X1[selected_cols1[0]][df['Cluster1'] == cluster]
                                if len(cluster_data) > 0:
                                    cluster_centers.append(cluster_data.mean())
                            cluster_centers.sort()
                            
                            for center in cluster_centers:
                                ax.axvline(center, color='black', linestyle='--', alpha=0.5)
                            
                            plt.title('第一个聚类的分布与簇分割')
                            plt.xlabel(selected_cols1[0])
                            plt.ylabel('密度')
                            plt.legend()
                            st.pyplot(fig)
                        
                        elif len(selected_cols1) == 2 or len(selected_cols1) > 2:
                            # 多维：选择可视化方式
                            vis_option1 = st.radio(
                                "选择可视化方式（第一个聚类）",
                                ["基于原始维度", "t-SNE降维到2维"],
                                index=0
                            )
                            
                            if vis_option1 == "t-SNE降维到2维":
                                # t-SNE 降维到 2 维
                                tsne = TSNE(n_components=2, random_state=42)
                                X_tsne = tsne.fit_transform(X_scaled1)
                                
                                fig, ax = plt.subplots(figsize=(10, 6))
                                sns.scatterplot(x=X_tsne[:, 0], y=X_tsne[:, 1], 
                                            hue=df['Cluster1'], palette='husl', ax=ax)
                                plt.title('第一个聚类的 t-SNE 降维分布')
                                plt.xlabel('t-SNE 维度 1')
                                plt.ylabel('t-SNE 维度 2')
                                st.pyplot(fig)
                            else:
                                # 基于原始维度（取前两个维度）
                                fig, ax = plt.subplots(figsize=(10, 6))
                                if len(selected_cols1) >= 2:
                                    sns.scatterplot(x=X1[selected_cols1[0]], y=X1[selected_cols1[1]], 
                                                hue=df['Cluster1'], palette='husl', ax=ax)
                                    plt.title('第一个聚类的前两个维度分布')
                                    plt.xlabel(selected_cols1[0])
                                    plt.ylabel(selected_cols1[1])
                                else:
                                    sns.histplot(X1[selected_cols1[0]], hue=df['Cluster1'], palette='husl', ax=ax)
                                    plt.title('第一个聚类的单维度分布')
                                    plt.xlabel(selected_cols1[0])
                                st.pyplot(fig)
                        
                        else:
                            # 多维：选择可视化方式
                            vis_option1 = st.radio(
                                "选择可视化方式（第一个聚类）",
                                ["基于原始维度", "t-SNE降维到2维"],
                                index=0
                            )
                            
                            if vis_option1 == "t-SNE降维到2维":
                                # t-SNE 降维到 2 维
                                tsne = TSNE(n_components=2, random_state=42)
                                X_tsne = tsne.fit_transform(X_scaled1)
                                
                                fig, ax = plt.subplots(figsize=(10, 6))
                                sns.scatterplot(x=X_tsne[:, 0], y=X_tsne[:, 1], 
                                            hue=df['Cluster1'], palette='husl', ax=ax)
                                plt.title('第一个聚类的 t-SNE 降维分布')
                                plt.xlabel('t-SNE 维度 1')
                                plt.ylabel('t-SNE 维度 2')
                                st.pyplot(fig)
                            else:
                                # 基于原始维度（取前两个维度）
                                fig, ax = plt.subplots(figsize=(10, 6))
                                sns.scatterplot(x=X1[selected_cols1[0]], y=X1[selected_cols1[1]], 
                                            hue=df['Cluster1'], palette='husl', ax=ax)
                                plt.title('第一个聚类的前两个维度分布')
                                plt.xlabel(selected_cols1[0])
                                plt.ylabel(selected_cols1[1])
                                st.pyplot(fig)
                        
                        cluster_col1 = 'Cluster1'
                    else:
                        st.warning("请至少选择一列用于聚类")
                else:
                    st.error("数据中没有数值型列，无法执行聚类")
            else:
                # 使用已有聚类列
                all_cols = df.columns.tolist()
                cluster_col1 = st.selectbox("选择第一个聚类列", all_cols)
                
                if cluster_col1:
                    df[cluster_col1] = df[cluster_col1].astype(str)  # 确保为分类变量
                    # 显示聚类结果
                    st.write("聚类结果分布:")
                    cluster_counts1 = df[cluster_col1].value_counts().sort_index()
                    st.write(cluster_counts1)
                    
                    # 绘制聚类分布图
                    fig, ax = plt.subplots(figsize=(10, 6))
                    cluster_counts1.plot(kind='bar', ax=ax)
                    plt.title('第一个聚类的簇分布')
                    plt.xlabel('簇编号')
                    plt.ylabel('样本数量')
                    st.pyplot(fig)
            
            # 第二个聚类
            st.subheader("第二个聚类")
            
            clustering_mode2 = st.radio(
                "选择第二个聚类模式:",
                ["使用模型聚类", "使用已有聚类列"],
                key="clustering_mode2"
            )
            
            if clustering_mode2 == "使用模型聚类":
                # 选择聚类算法
                clustering_algo2 = st.selectbox("选择聚类算法", ["K-means", "DBSCAN"], key="clustering_algo2")
                # 添加簇命名风格选择
                naming_style2 = st.selectbox(
                    "选择簇命名风格",
                    ["字母", "数字", "中文"],
                    index=0,  # 默认选择字母
                    key="naming_style2"
                )
                # 选择用于聚类的列
                numeric_cols = df.select_dtypes(include=['int64', 'float64']).columns.tolist()
                if len(numeric_cols) > 0:
                    selected_cols2 = st.multiselect(
                        "选择用于第二个聚类的列",
                        numeric_cols,
                        default=numeric_cols[-max(0,min(3, len(numeric_cols))):]
                    )
                    
                    if selected_cols2:
                        X2 = df[selected_cols2]
                        # 标准化数据
                        scaler = StandardScaler()
                        X_scaled2 = scaler.fit_transform(X2)
                        
                        # 多维数据降维选项
                        if len(selected_cols2) >= 2:
                            input_option2 = st.radio(
                                "选择聚类输入数据（第二个聚类）",
                                ["原始数据", "t-SNE降维到2维"],
                                index=0
                            )
                            if input_option2 == "t-SNE降维到2维":
                                tsne = TSNE(n_components=2, random_state=42)
                                X_scaled2 = tsne.fit_transform(X_scaled2)
                        
                        if clustering_algo2 == "K-means":
                            # 设置 K 值
                            k_value2 = st.slider("选择 K-means 的 K 值（簇的数量）", 2, 10, 3, key="k_value2")
                            
                            # 执行 K-means
                            kmeans2 = KMeans(n_clusters=k_value2, random_state=42, n_init=10)
                            # 替换为:
                            labels = kmeans2.fit_predict(X_scaled2)
                            df['Cluster2'] = [convert_cluster_label(str(label), naming_style2) for label in labels]

                            # 计算轮廓系数
                            silhouette_avg = silhouette_score(X_scaled2, df['Cluster2'])
                            st.write(f"轮廓系数 (越接近1越好): {silhouette_avg:.4f}")
                        
                        else:  # DBSCAN
                            # 设置 DBSCAN 参数并添加通俗说明
                            st.markdown("""
                            **DBSCAN 参数说明**:
                            - **eps**: 点的邻域半径，决定多远的点被认为是“邻居”。值越小，簇越密集；值越大，簇越分散。
                            - **min_samples**: 一个簇所需的最小点数（包括核心点本身）。值越大，簇需要更多点才能形成；值越小，可能生成更多小簇。
                            """)
                            eps2 = st.slider("选择 DBSCAN 的 eps 参数", 0.1, 2.0, 0.5, step=0.1, key="eps2")
                            min_samples2 = st.slider("选择 DBSCAN 的 min_samples 参数", 2, 20, 5, key="min_samples2")
                            
                            # 执行 DBSCAN
                            dbscan2 = DBSCAN(eps=eps2, min_samples=min_samples2)
                            labels = dbscan2.fit_predict(X_scaled2)
                            
                            # 处理噪声点 (-1)
                            if -1 in labels:
                                # 找到非噪声点
                                non_noise_mask = labels != -1
                                X_non_noise = X_scaled2[non_noise_mask]
                                labels_non_noise = labels[non_noise_mask]
                                
                                # 对噪声点使用 KNN 分配最近的簇
                                noise_mask = labels == -1
                                X_noise = X_scaled2[noise_mask]
                                
                                if len(X_non_noise) > 0 and len(X_noise) > 0:
                                    knn = NearestNeighbors(n_neighbors=1)
                                    knn.fit(X_non_noise)
                                    _, indices = knn.kneighbors(X_noise)
                                    labels[noise_mask] = labels_non_noise[indices.flatten()]
                            
                            df['Cluster2'] = [convert_cluster_label(str(label), naming_style2) for label in labels]
                            
                            # 检查是否有有效簇
                            if len(np.unique(df['Cluster2'])) > 1:
                                silhouette_avg = silhouette_score(X_scaled2, df['Cluster2'])
                                st.write(f"轮廓系数 (越接近1越好): {silhouette_avg:.4f}")
                            else:
                                st.warning("DBSCAN 未能生成多个有效簇，请调整 eps 或 min_samples 参数")
                        
                        # 显示聚类结果
                        st.write("聚类结果分布:")
                        cluster_counts2 = df['Cluster2'].value_counts().sort_index()
                        st.write(cluster_counts2)
                        
                        # 可视化聚类结果
                        if len(selected_cols2) == 1:
                            # 一维：分布图 + 簇分割线
                            fig, ax = plt.subplots(figsize=(10, 6))
                            unique_clusters = sorted(df['Cluster2'].unique())
                            colors = sns.color_palette("husl", len(unique_clusters))
                            
                            for cluster, color in zip(unique_clusters, colors):
                                cluster_data = X2[selected_cols2[0]][df['Cluster2'] == cluster]
                                sns.histplot(cluster_data, kde=True, label=f'簇 {cluster}', 
                                        stat='density', alpha=0.4, color=color, ax=ax)
                            
                            # 添加簇分割线（基于簇中心）
                            cluster_centers = []
                            for cluster in unique_clusters:
                                cluster_data = X2[selected_cols2[0]][df['Cluster2'] == cluster]
                                if len(cluster_data) > 0:
                                    cluster_centers.append(cluster_data.mean())
                            cluster_centers.sort()
                            
                            for center in cluster_centers:
                                ax.axvline(center, color='black', linestyle='--', alpha=0.5)
                            
                            plt.title('第二个聚类的分布与簇分割')
                            plt.xlabel(selected_cols2[0])
                            plt.ylabel('密度')
                            plt.legend()
                            st.pyplot(fig)
                        
                        elif len(selected_cols2) == 2 or len(selected_cols2) > 2:
                            # 多维：选择可视化方式
                            vis_option2 = st.radio(
                                "选择可视化方式（第二个聚类）",
                                ["基于原始维度", "t-SNE降维到2维"],
                                index=0
                            )
                            
                            if vis_option2 == "t-SNE降维到2维":
                                # t-SNE 降维到 2 维
                                tsne = TSNE(n_components=2, random_state=42)
                                X_tsne = tsne.fit_transform(X_scaled2)
                                
                                fig, ax = plt.subplots(figsize=(10, 6))
                                sns.scatterplot(x=X_tsne[:, 0], y=X_tsne[:, 1], 
                                            hue=df['Cluster2'], palette='husl', ax=ax)
                                plt.title('第二个聚类的 t-SNE 降维分布')
                                plt.xlabel('t-SNE 维度 1')
                                plt.ylabel('t-SNE 维度 2')
                                st.pyplot(fig)
                            else:
                                # 基于原始维度（取前两个维度）
                                fig, ax = plt.subplots(figsize=(10, 6))
                                if len(selected_cols2) >= 2:
                                    sns.scatterplot(x=X2[selected_cols2[0]], y=X2[selected_cols2[1]], 
                                                hue=df['Cluster2'], palette='husl', ax=ax)
                                    plt.title('第二个聚类的前两个维度分布')
                                    plt.xlabel(selected_cols2[0])
                                    plt.ylabel(selected_cols2[1])
                                else:
                                    sns.histplot(X2[selected_cols2[0]], hue=df['Cluster2'], palette='husl', ax=ax)
                                    plt.title('第二个聚类的单维度分布')
                                    plt.xlabel(selected_cols2[0])
                                st.pyplot(fig)
                        
                        else:
                            # 多维：选择可视化方式
                            vis_option2 = st.radio(
                                "选择可视化方式（第二个聚类）",
                                ["基于原始维度", "t-SNE降维到2维"],
                                index=0
                            )
                            
                            if vis_option2 == "t-SNE降维到2维":
                                # t-SNE 降维到 2 维
                                tsne = TSNE(n_components=2, random_state=42)
                                X_tsne = tsne.fit_transform(X_scaled2)
                                
                                fig, ax = plt.subplots(figsize=(10, 6))
                                sns.scatterplot(x=X_tsne[:, 0], y=X_tsne[:, 1], 
                                            hue=df['Cluster2'], palette='husl', ax=ax)
                                plt.title('第二个聚类的 t-SNE 降维分布')
                                plt.xlabel('t-SNE 维度 1')
                                plt.ylabel('t-SNE 维度 2')
                                st.pyplot(fig)
                            else:
                                # 基于原始维度（取前两个维度）
                                fig, ax = plt.subplots(figsize=(10, 6))
                                sns.scatterplot(x=X2[selected_cols2[0]], y=X2[selected_cols2[1]], 
                                            hue=df['Cluster2'], palette='husl', ax=ax)
                                plt.title('第二个聚类的前两个维度分布')
                                plt.xlabel(selected_cols2[0])
                                plt.ylabel(selected_cols2[1])
                                st.pyplot(fig)
                        
                        cluster_col2 = 'Cluster2'
                    else:
                        st.warning("请至少选择一列用于聚类")
                else:
                    st.error("数据中没有数值型列，无法执行聚类")
            else:
                # 使用已有聚类列
                all_cols = df.columns.tolist()
                if cluster_col1 in all_cols:
                    all_cols.remove(cluster_col1)
                
                cluster_col2 = st.selectbox("选择第二个聚类列", all_cols)
                
                if cluster_col2:
                    df[cluster_col2] = df[cluster_col2].astype(str)  # 确保为分类变量
                    # 显示聚类结果
                    st.write("聚类结果分布:")
                    cluster_counts2 = df[cluster_col2].value_counts().sort_index()
                    st.write(cluster_counts2)
                    
                    # 绘制聚类分布图
                    fig, ax = plt.subplots(figsize=(10, 6))
                    cluster_counts2.plot(kind='bar', ax=ax)
                    plt.title('第二个聚类的簇分布')
                    plt.xlabel('簇编号')
                    plt.ylabel('样本数量')
                    st.pyplot(fig)
            
            # 热力图分析
            if cluster_col1 and cluster_col2:
                st.subheader("热力图分析")
                
                # 创建交叉表
                crosstab = pd.crosstab(df[cluster_col1], df[cluster_col2])
                
                # 计算期望频率
                chi2, p, dof, expected = chi2_contingency(crosstab)
                
                # 计算调整后残差
                observed = crosstab.values
                expected = expected.reshape(observed.shape)
                
                # 计算残差
                residuals = observed - expected
                
                # 计算调整后残差
                n = observed.sum()
                row_sums = observed.sum(axis=1).reshape(-1, 1)
                col_sums = observed.sum(axis=0).reshape(1, -1)
                
                adj_residuals = residuals / np.sqrt(
                    expected * (1 - row_sums / n) * (1 - col_sums / n)
                )
                
                # 创建调整后残差的 DataFrame
                adj_residuals_df = pd.DataFrame(
                    adj_residuals,
                    index=crosstab.index,
                    columns=crosstab.columns
                )
                
                # 显示交叉表
                st.write("交叉表 (观察值):")
                st.dataframe(crosstab)
                
                # 绘制热力图（观察值）
                st.write("观察值热力图:")
                fig, ax = plt.subplots(figsize=(12, 8))
                sns.heatmap(crosstab, annot=True, fmt="d", cmap="YlGnBu", ax=ax)
                plt.title(f'{cluster_col1} 与 {cluster_col2} 的交叉热力图')
                st.pyplot(fig)
                
                # 绘制调整后残差热力图
                st.write("调整后残差热力图:")
                st.write("(值 > 1.96 或 < -1.96 表示在95%置信度下统计显著)")
                fig, ax = plt.subplots(figsize=(12, 8))
                sns.heatmap(adj_residuals_df, annot=True, fmt=".2f", cmap="coolwarm", center=0, ax=ax)
                plt.title(f'{cluster_col1} 与 {cluster_col2} 的调整后残差热力图')
                st.pyplot(fig)
                
                # 显示卡方检验结果
                st.write(f"卡方统计量: {chi2:.4f}, p值: {p:.4f}")
                if p < 0.05:
                    st.write("两个聚类之间存在显著关联 (p < 0.05)")
                else:
                    st.write("两个聚类之间不存在显著关联 (p >= 0.05)")
                
                # 查找最显著的组合
                st.subheader("最显著的聚类组合:")
                flat_residuals = adj_residuals_df.abs().stack()  # 使用 stack() 替代 unstack()
                top_significant = flat_residuals.sort_values(ascending=False).head(5)
                
                for idx, value in top_significant.items():
                    cluster1, cluster2 = idx
                    try:
                        observed_val = crosstab.loc[cluster1, cluster2]
                        expected_val = expected[crosstab.index.get_loc(cluster1), crosstab.columns.get_loc(cluster2)]
                        direction = "高于" if observed_val > expected_val else "低于"
                        
                        st.write(f"聚类组合 ({cluster_col1}={cluster1}, {cluster_col2}={cluster2}):")
                        st.write(f"  - 调整后残差: {adj_residuals_df.loc[cluster1, cluster2]:.4f}")
                        st.write(f"  - 观察计数: {observed_val} ({direction}期望)")
                        st.write(f"  - 期望计数: {expected_val:.2f}")
                    except KeyError:
                        st.warning(f"聚类组合 ({cluster_col1}={cluster1}, {cluster_col2}={cluster2})无效：索引不存在")
                
                # 提供下载功能
                st.subheader("下载分析结果")
                
                # 准备数据
                output = io.BytesIO()
                with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                    df.to_excel(writer, sheet_name='原始数据与聚类', index=False)
                    crosstab.to_excel(writer, sheet_name='交叉表')
                    adj_residuals_df.to_excel(writer, sheet_name='调整后残差')
                    
                    # 添加最显著组合的表格
                    sig_data = []
                    for idx, value in top_significant.items():
                        cluster1, cluster2 = idx
                        try:
                            observed_val = crosstab.loc[cluster1, cluster2]
                            expected_val = expected[crosstab.index.get_loc(cluster1), crosstab.columns.get_loc(cluster2)]
                            sig_data.append({
                                f'{cluster_col1}': cluster1,
                                f'{cluster_col2}': cluster2,
                                '调整后残差': adj_residuals_df.loc[cluster1, cluster2],
                                '观察计数': observed_val,
                                '期望计数': expected_val,
                                '是否显著': abs(adj_residuals_df.loc[cluster1, cluster2]) > 1.96
                            })
                        except KeyError:
                            continue
                    
                    pd.DataFrame(sig_data).to_excel(writer, sheet_name='显著组合')
                
                output.seek(0)
                
                # 提供下载链接
                b64 = base64.b64encode(output.read()).decode()
                href = f'<a href="data:application/vnd.openxmlformats-officedocument.spreadsheetml.sheet;base64,{b64}" download="聚类分析结果.xlsx">下载Excel分析结果</a>'
                st.markdown(href, unsafe_allow_html=True)
                
        except Exception as e:
            import traceback
            st.error(f"处理文件时出错: {str(e)}")
            st.error(f"错误详情: {traceback.format_exc()}")
    else:
        st.info("请上传 CSV 或 Excel 文件以开始分析")

# Modify the main() function to add the Sales Forecasting option
def main():
    st.set_page_config(layout="wide")

    st.markdown(
        """
        <h3 style='color: #800080; font-weight: bold; text-align: center;'>
            🎵 随时随地聆听精彩音频！🎵
        </h3>
        <p style='color: #800080; text-align: center;'>
            探索 app 的同时，享受独家医疗洞察音频（信息来源于网络），提升您的使用体验！
        </p>
        """,
        unsafe_allow_html=True
    )
    
    # Audio selection dropdown
    audio_folder = "audio_folder"
    default_audio = "sample.mp3"
    audio_files = []
    
    # Check if audio_folder exists and list .mp3 files
    if os.path.exists(audio_folder):
        audio_files = [f for f in os.listdir(audio_folder) if f.endswith(".mp3")]
    
    # Add default audio to the list if it exists in the folder
    if audio_files and default_audio in audio_files:
        default_index = audio_files.index(default_audio)
    else:
        default_index = 0
        if not audio_files:
            audio_files = ["无音频文件可用"]
    
    # Audio dropdown menu
    selected_audio = st.selectbox(
        "选择音频",
        audio_files,
        index=default_index,
        key="audio_select"
    )
    
    # Play selected audio
    if selected_audio != "无音频文件可用":
        audio_file_path = os.path.join(audio_folder, selected_audio)
        if os.path.exists(audio_file_path):
            st.audio(audio_file_path, format="audio/mp3", start_time=0)
            # st.markdown(
            #     "<p style='color: #00FF00; text-align: center; font-weight: bold;'>正在播放中...</p>",
            #     unsafe_allow_html=True
            # )
        else:
            st.error(f"音频文件未找到: {audio_file_path}")
    else:
        st.error("音频文件夹为空或不存在，请检查 audio_folder/ 目录")
    
    # Create the page selection in sidebar
    page = st.sidebar.radio("选择功能", ["Medical Insights Copilot", "Spreadsheet Analysis", "Sales Forecasting","Cluster Analysis"])
    
    if page == "Medical Insights Copilot":  
        # model_choice, client = setup_client(model_choice = 'gemini-2.0-flash')
        model_choice, client = setup_client()
        setup_layout(
            topics, diseases, institutions, departments, persons,
            primary_topics_list, primary_diseases_list,
            generate_tag, generate_diseases_tag, rewrite,
            prob_identy, generate_structure_data,
            model_choice, client
        )
    elif page == "Spreadsheet Analysis":
        setup_spreadsheet_analysis()
    elif page == "Sales Forecasting":
        setup_sales_forecasting()
    elif page == "Cluster Analysis":
        setup_clustering_analysis()


if __name__ == "__main__":
    main()
